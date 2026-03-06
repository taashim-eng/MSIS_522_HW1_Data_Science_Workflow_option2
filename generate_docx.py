"""
Generate research_writeup.docx for MSIS 522 HW1 Alternative Credit Scoring
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    hs.font.name = 'Calibri'

# ── Helper functions ──
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacer

def add_code_block(doc, code, language=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Light gray shading via XML
    shading = run._element.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:fill'): 'F4F4F4',
    })
    run._element.get_or_add_rPr().append(shading)

def add_bullet(doc, text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Alternative Credit Scoring\nDataset Research & HW1 Alignment')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('\nMSIS 522 B — Advanced Machine Learning\nHomework 1 — The Complete Data Science Workflow\n\n2026-03-05')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This report identifies, evaluates, and documents six candidate tabular datasets '
    'suitable for building explainable alternative credit-scoring models. The datasets span '
    'classical credit-risk benchmarks (German Credit, Taiwan Default, FICO HELOC), Kaggle '
    'competition data (Give Me Some Credit, Home Credit Default Risk), and peer-to-peer '
    'lending records (LendingClub). Each dataset supports binary classification tasks aligned '
    'with MSIS 522 HW1: predicting whether a borrower will default.'
)

doc.add_heading('Top 2 Recommendations for HW1', level=3)
add_table(doc,
    ['Rank', 'Dataset', 'Why'],
    [
        ['#1', 'Taiwan Default of Credit Card Clients',
         '30K rows, single CSV, clean binary target, demographic features for fairness analysis, well-documented in academic literature'],
        ['#2', 'Statlog German Credit',
         '1K-row classic benchmark, interpretable features, built-in fairness concerns (personal status/sex, foreign worker), fast prototyping'],
    ])

doc.add_paragraph(
    'Both are directly downloadable without credentials and satisfy all HW1 requirements: '
    '≥200 rows, clear target variable, tabular format, and real-world provenance.'
)

# ══════════════════════════════════════════════════════════════
# 2. DATASET INVENTORY
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('2. Dataset Inventory', level=1)

# ── Dataset 1: German Credit ──
doc.add_heading('Dataset 1: Statlog (German Credit Data)', level=2)
add_table(doc,
    ['Attribute', 'Detail'],
    [
        ['Title', 'Statlog (German Credit Data)'],
        ['Summary', 'Classify 1,000 bank customers as good or bad credit risks using 20 financial/demographic features.'],
        ['Source', 'UCI Machine Learning Repository'],
        ['Direct URL', 'https://archive.ics.uci.edu/ml/machine-learning-databases/statlog/german/german.data'],
        ['Files', 'german.data (79,793 B), german.data-numeric (102,000 B), german.doc (4,679 B)'],
        ['Format', 'Space-separated text, no header'],
        ['Shape', '1,000 rows × 21 columns'],
        ['Target', 'Class — 1 = Good, 2 = Bad'],
        ['Key Fields', 'Status_checking_acct, Duration, Credit_history, Purpose, Credit_amount, Savings_acct, Age, Personal_status_sex, Foreign_worker'],
        ['License', 'CC BY 4.0'],
        ['Missing Values', '0'],
        ['PII Risk', 'None — fully anonymized categorical codes'],
    ])

doc.add_heading('Sample Extraction', level=3)
add_code_block(doc, '''import pandas as pd
cols = ['Status_checking_acct','Duration','Credit_history','Purpose','Credit_amount',
        'Savings_acct','Present_employment','Installment_rate','Personal_status_sex',
        'Other_debtors','Residence_since','Property','Age','Other_installment_plans',
        'Housing','Existing_credits','Job','Num_dependents','Telephone','Foreign_worker','Class']
df = pd.read_csv('german-credit/raw/german.data', sep=' ', header=None, names=cols)
sample = df.sample(n=200, random_state=42)
sample.to_csv('german-credit/sample.csv', index=False)''')

doc.add_heading('HW1 Mapping', level=3)
doc.add_paragraph(
    'This dataset supports all five models (Logistic Regression baseline, Decision Tree, '
    'Random Forest, XGBoost, MLP) with its mix of categorical and numeric features. The '
    'Personal_status_sex and Foreign_worker columns make it ideal for Tab 5 Ethics & Bias '
    'analysis (disparate impact by gender, nationality). Its small size allows rapid '
    'GridSearchCV iteration. SHAP waterfall plots are straightforward since the feature space '
    'is interpretable. For Tab 4, expose sliders for Duration, Credit_amount, Age, and dropdowns '
    'for Purpose, Status_checking_acct.'
)

# ── Dataset 2: Taiwan Credit ──
doc.add_page_break()
doc.add_heading('Dataset 2: Default of Credit Card Clients (Taiwan)', level=2)
add_table(doc,
    ['Attribute', 'Detail'],
    [
        ['Title', 'Default of Credit Card Clients'],
        ['Summary', 'Predict whether 30,000 Taiwan credit-card holders will default on payment next month.'],
        ['Source', 'UCI Machine Learning Repository'],
        ['Direct URL', 'https://archive.ics.uci.edu/static/public/350/default+of+credit+card+clients.zip'],
        ['Files', 'default of credit card clients.xls (5,539,328 B)'],
        ['Format', 'Excel (.xls), header on row 2'],
        ['Shape', '30,000 rows × 25 columns (incl. ID)'],
        ['Target', 'default payment next month — 0 = No, 1 = Yes'],
        ['Key Fields', 'LIMIT_BAL, SEX, EDUCATION, MARRIAGE, AGE, PAY_0–PAY_6, BILL_AMT1–BILL_AMT6, PAY_AMT1–PAY_AMT6'],
        ['License', 'CC BY 4.0'],
        ['Missing Values', '0'],
        ['PII Risk', 'None — anonymized IDs, no names/addresses'],
    ])

doc.add_heading('Sample Extraction', level=3)
add_code_block(doc, '''import pandas as pd
df = pd.read_excel('taiwan-credit/raw/default of credit card clients.xls', header=1)
sample = df.sample(n=200, random_state=42)
sample.to_csv('taiwan-credit/sample.csv', index=False)''')

doc.add_heading('HW1 Mapping', level=3)
doc.add_paragraph(
    'The recommended primary dataset for HW1. At 30K rows, it provides robust training for all '
    'five models. The 6-month payment history (PAY_0–PAY_6) and billing amounts offer rich '
    'numeric features for SHAP interpretation. SEX, EDUCATION, MARRIAGE, and AGE enable fairness '
    'auditing (Tab 5). For Tab 4 interactive predictions, expose sliders for LIMIT_BAL, AGE, and '
    'PAY_0 (most recent payment status), with dropdowns for SEX, EDUCATION, MARRIAGE; use '
    'mean/mode for remaining features. Class imbalance (~22% default) requires handling via SMOTE '
    'or class_weight="balanced".'
)

# ── Dataset 3: FICO HELOC ──
doc.add_page_break()
doc.add_heading('Dataset 3: FICO HELOC (Home Equity Line of Credit)', level=2)
add_table(doc,
    ['Attribute', 'Detail'],
    [
        ['Title', 'FICO Explainable ML Challenge — HELOC Dataset'],
        ['Summary', 'Predict "good/bad" repayment performance on 10,459 HELOC applications using 23 credit-bureau features.'],
        ['Source', 'FICO Community / GitHub mirror'],
        ['Direct URL', 'https://raw.githubusercontent.com/benoitparis/explainable-challenge/master/heloc_dataset_v1.csv'],
        ['Files', 'heloc_dataset_v1.csv (678,479 B)'],
        ['Format', 'CSV with header'],
        ['Shape', '10,459 rows × 24 columns'],
        ['Target', 'RiskPerformance — "Good" or "Bad"'],
        ['Key Fields', 'ExternalRiskEstimate, MSinceOldestTradeOpen, NumSatisfactoryTrades, PercentTradesNeverDelq, NetFractionRevolvingBurden, NumInqLast6M'],
        ['License', 'Research use (FICO Challenge terms)'],
        ['Missing Values', '0 (sentinel values -7, -8, -9 encode special conditions)'],
        ['PII Risk', 'None — all credit-bureau aggregates'],
    ])

doc.add_heading('HW1 Mapping', level=3)
doc.add_paragraph(
    'Designed specifically for explainable machine learning. All 23 features are numeric '
    'credit-bureau aggregates, making SHAP analysis especially clear. The sentinel values '
    '(-7, -8, -9) require pre-processing (replace with NaN or encode). No demographic features, '
    'so fairness analysis must focus on algorithmic bias rather than protected attributes. '
    'For Tab 4, expose sliders for ExternalRiskEstimate, NumSatisfactoryTrades, '
    'PercentTradesNeverDelq, NetFractionRevolvingBurden; average remaining features.'
)

# ── Dataset 4: Give Me Some Credit ──
doc.add_page_break()
doc.add_heading('Dataset 4: Give Me Some Credit (Kaggle)', level=2)
add_table(doc,
    ['Attribute', 'Detail'],
    [
        ['Title', 'Give Me Some Credit'],
        ['Summary', 'Predict financial distress within 2 years for 150,000 borrowers using 10 financial features.'],
        ['Source', 'Kaggle Competition'],
        ['Kaggle Slug', 'competitions/GiveMeSomeCredit'],
        ['Download Cmd', 'kaggle competitions download -c GiveMeSomeCredit -p give-me-some-credit/raw/'],
        ['Credentials', 'Set KAGGLE_USERNAME and KAGGLE_KEY environment variables'],
        ['Shape', '~150,000 rows × 11 columns'],
        ['Target', 'SeriousDlqin2yrs — 0 = No, 1 = Yes'],
        ['Key Fields', 'RevolvingUtilizationOfUnsecuredLines, age, DebtRatio, MonthlyIncome, NumberOfDependents'],
        ['License', 'Kaggle Competition Rules'],
        ['Missing Values', '~29,000 in MonthlyIncome, ~3,900 in NumberOfDependents'],
        ['PII Risk', 'None — anonymized'],
    ])

doc.add_heading('HW1 Mapping', level=3)
doc.add_paragraph(
    'Large training set for robust model comparison. Missing values in MonthlyIncome and '
    'NumberOfDependents demonstrate imputation techniques. The age field supports basic '
    'fairness checks. Target is imbalanced (~6.7% positive), motivating SMOTE or class weighting. '
    'Good candidate for demonstrating the full GridSearchCV pipeline due to ample data.'
)

# ── Dataset 5: Home Credit ──
doc.add_heading('Dataset 5: Home Credit Default Risk (Kaggle)', level=2)
add_table(doc,
    ['Attribute', 'Detail'],
    [
        ['Title', 'Home Credit Default Risk'],
        ['Summary', 'Predict loan repayment ability for 307K+ applicants using multi-table financial data.'],
        ['Source', 'Kaggle Competition'],
        ['Kaggle Slug', 'competitions/home-credit-default-risk'],
        ['Download Cmd', 'kaggle competitions download -c home-credit-default-risk -p home-credit-default-risk/raw/'],
        ['Shape', '307,511 rows × 122 columns (main table)'],
        ['Target', 'TARGET — 0 = Repaid, 1 = Default'],
        ['Key Fields', 'AMT_INCOME_TOTAL, AMT_CREDIT, AMT_ANNUITY, CODE_GENDER, DAYS_BIRTH, EXT_SOURCE_1/2/3'],
        ['License', 'Kaggle Competition Rules'],
        ['Missing Values', 'Significant — 60+ columns have >50% missing'],
        ['PII Risk', 'None — anonymized with encoded categories'],
    ])

doc.add_heading('HW1 Mapping', level=3)
doc.add_paragraph(
    'The richest dataset with relational tables supporting advanced feature engineering. '
    'However, its 122 columns and multi-table structure may be overly complex for HW1 — '
    'better suited for a semester project. The CODE_GENDER field enables fairness analysis. '
    'Has severe class imbalance (~8% default).'
)

# ── Dataset 6: LendingClub ──
doc.add_heading('Dataset 6: LendingClub Loan Data (Kaggle)', level=2)
add_table(doc,
    ['Attribute', 'Detail'],
    [
        ['Title', 'Lending Club Loans 2007-2020'],
        ['Summary', 'Peer-to-peer lending records for ~2.9M loans with 151 features.'],
        ['Source', 'Kaggle'],
        ['Kaggle Slug', 'datasets/ethon0426/lending-club-20072020q1'],
        ['Download Cmd', 'kaggle datasets download -d ethon0426/lending-club-20072020q1 -p lendingclub-loans/raw/'],
        ['Shape', '~2,900,000 rows × 151 columns'],
        ['Target', 'loan_status — map "Charged Off" = 1, "Fully Paid" = 0'],
        ['Key Fields', 'loan_amnt, int_rate, annual_inc, dti, delinq_2yrs, fico_range_low, purpose, addr_state'],
        ['License', 'CC0 Public Domain'],
        ['Missing Values', 'Substantial across many columns'],
        ['PII Risk', 'Low — addr_state and zip_code (3-digit) may be quasi-identifiers'],
    ])

# ══════════════════════════════════════════════════════════════
# 3. TOP 2 RECOMMENDATION
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('3. Top 2 Recommendation & Justification', level=1)

doc.add_heading('Primary: Taiwan Default of Credit Card Clients', level=2)
add_bullet(doc, '30,000 rows — large enough for meaningful train/test splits and GridSearchCV, small enough for fast MLP training', 'Right size: ')
add_bullet(doc, 'Zero missing values, single file, straightforward Excel import', 'Clean: ')
add_bullet(doc, 'Clear target (default payment next month)', 'Binary classification: ')
add_bullet(doc, 'SEX, EDUCATION, MARRIAGE, AGE enable Tab 5 disparate-impact and equalized-odds analysis', 'Fairness-ready: ')
add_bullet(doc, 'Mix of categorical (encoded as integers) and continuous features produces interpretable plots', 'SHAP-friendly: ')
add_bullet(doc, '~22% default rate — enough to motivate SMOTE/class-weighting, not so extreme as to collapse metrics', 'Class imbalance: ')
add_bullet(doc, 'Published in Yeh & Lien (2009), widely used in ML coursework', 'Academic pedigree: ')

doc.add_heading('Secondary: Statlog German Credit', level=2)
add_bullet(doc, '1,000 rows — entire pipeline runs in seconds', 'Fast prototyping: ')
add_bullet(doc, '20 features with clear business meaning', 'Interpretable: ')
add_bullet(doc, 'Personal_status_sex, Foreign_worker, Age are textbook fairness-audit examples', 'Ethics showcase: ')
add_bullet(doc, 'One of the most-cited credit datasets in ML literature', 'Proven benchmark: ')
add_bullet(doc, 'If used alongside Taiwan Credit, demonstrates model behavior on small vs. medium datasets', 'Complementary: ')

# ══════════════════════════════════════════════════════════════
# 4. EDA PLAN
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('4. EDA Plan (Top 2 Picks)', level=1)

doc.add_heading('Taiwan Default of Credit Card Clients', level=2)
add_table(doc,
    ['EDA Step', 'Detail'],
    [
        ['Distributions', 'Histograms for LIMIT_BAL, AGE, BILL_AMT1–BILL_AMT6'],
        ['Target Balance', 'Bar chart of default payment next month (0 vs. 1)'],
        ['Correlation', 'Heatmap of all numeric features (expect high correlation among BILL_AMT and PAY_AMT series)'],
        ['Cross-tabs', 'Default rate by SEX, EDUCATION, MARRIAGE'],
        ['Outlier Detection', 'Box plots for LIMIT_BAL, BILL_AMT1, PAY_AMT1'],
        ['Payment Patterns', 'Stacked bar chart of PAY_0–PAY_6 distributions'],
    ])

doc.add_heading('Statlog German Credit', level=2)
add_table(doc,
    ['EDA Step', 'Detail'],
    [
        ['Distributions', 'Histograms for Duration, Credit_amount, Age'],
        ['Target Balance', 'Bar chart of Class (1=Good vs. 2=Bad; note: 70/30 split)'],
        ['Cross-tabs', 'Default rate by Personal_status_sex, Foreign_worker, Purpose'],
        ['Correlation', 'Heatmap of numeric features (after encoding categoricals)'],
        ['Feature Importance', 'Preliminary Random Forest importance plot'],
    ])

# ══════════════════════════════════════════════════════════════
# 5. FEATURE ENGINEERING
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Feature Engineering Suggestions', level=1)

doc.add_heading('Taiwan Credit', level=2)
add_bullet(doc, 'avg_delay = mean(PAY_0..PAY_6), max_delay, num_months_delayed', 'Payment behavior aggregates: ')
add_bullet(doc, 'BILL_AMT1 / LIMIT_BAL', 'Utilization ratio: ')
add_bullet(doc, 'PAY_AMT1 / BILL_AMT1 (how much of the bill was paid)', 'Payment ratio: ')
add_bullet(doc, 'Slope of BILL_AMT1..BILL_AMT6 over time (increasing debt trend)', 'Trend features: ')
add_bullet(doc, 'Discretize AGE into brackets for fairness subgroup analysis', 'Age bins: ')

doc.add_heading('German Credit', level=2)
add_bullet(doc, 'One-hot or ordinal encoding for Status_checking_acct, Credit_history, Purpose, etc.', 'Encode categoricals: ')
add_bullet(doc, 'Credit_amount × Duration, Age × Present_employment', 'Interaction features: ')
add_bullet(doc, 'has_telephone, is_foreign_worker', 'Binary flags: ')
add_bullet(doc, 'Combine Status_checking_acct + Savings_acct into liquidity score', 'Risk score proxy: ')

# ══════════════════════════════════════════════════════════════
# 6. BASELINE MODEL PLAN
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('6. Baseline Model Plan', level=1)

doc.add_heading('Train/Test Split', level=3)
add_code_block(doc, '''from sklearn.model_selection import train_test_split
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.3, random_state=42, stratify=y
)''')

doc.add_heading('Model Pipeline', level=3)
add_table(doc,
    ['#', 'Model', 'Implementation', 'Hyperparameter Tuning'],
    [
        ['1', 'Logistic Regression', 'sklearn.linear_model.LogisticRegression', 'Baseline — default params'],
        ['2', 'Decision Tree', 'sklearn.tree.DecisionTreeClassifier', 'GridSearchCV(5-fold): max_depth=[3,5,7,10], min_samples_leaf=[5,10,20,50]'],
        ['3', 'Random Forest', 'sklearn.ensemble.RandomForestClassifier', 'GridSearchCV(5-fold): n_estimators=[50,100,200], max_depth=[3,5,8]'],
        ['4', 'XGBoost', 'xgboost.XGBClassifier', 'GridSearchCV(5-fold): learning_rate=[0.01,0.1,0.3], n_estimators=[50,100,200], max_depth=[3,5,7]'],
        ['5', 'MLP', 'keras.Sequential', '≥2 hidden layers (128 units, ReLU), output sigmoid, binary crossentropy'],
    ])

doc.add_heading('Evaluation Metrics', level=3)
add_bullet(doc, ', '.join(['Accuracy', 'Precision', 'Recall', 'F1-Score', 'ROC-AUC']))
add_bullet(doc, 'Confusion matrix for each model')
add_bullet(doc, 'ROC curve comparison plot')

# ══════════════════════════════════════════════════════════════
# 7. FAIRNESS & AUDIT
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Fairness & Audit Checklist', level=1)

doc.add_heading('Disparate Impact Analysis', level=3)
add_code_block(doc, '''# For Taiwan Credit — check default predictions by SEX
from sklearn.metrics import confusion_matrix
male_preds = model.predict(X_test[X_test['SEX'] == 1])
female_preds = model.predict(X_test[X_test['SEX'] == 2])
# Disparate impact ratio = P(favorable|unprivileged) / P(favorable|privileged)''')

doc.add_heading('Checklist', level=3)
add_bullet(doc, 'Check class distribution across protected attributes (SEX, AGE, EDUCATION, MARRIAGE)', 'Data bias: ')
add_bullet(doc, 'Compute DI ratio for each model; flag if < 0.8 (four-fifths rule)', 'Disparate impact: ')
add_bullet(doc, 'Compare TPR and FPR across subgroups', 'Equalized odds: ')
add_bullet(doc, 'Check if predicted probabilities are well-calibrated per subgroup', 'Calibration: ')
add_bullet(doc, 'Use SHAP to check if protected attributes have outsized importance', 'Feature influence: ')
add_bullet(doc, 'If bias detected: remove protected features, reweight, or post-processing threshold adjustment', 'Mitigation: ')
add_bullet(doc, 'Record all findings in Tab 5 of Streamlit app', 'Documentation: ')

# ══════════════════════════════════════════════════════════════
# 8. STREAMLIT WIREFRAME
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('8. Streamlit App Wireframe', level=1)

doc.add_heading('Tab 1: Executive Summary', level=2)
add_bullet(doc, 'Problem statement and dataset description')
add_bullet(doc, 'Key findings in bullet points')
add_bullet(doc, 'Best model name and its AUC-ROC score')

doc.add_heading('Tab 2: Descriptive Analytics', level=2)
add_bullet(doc, 'st.dataframe(df.head()) — interactive data preview')
add_bullet(doc, 'st.write(df.describe()) — summary statistics')
add_bullet(doc, 'Plotly/Matplotlib histograms for key features')
add_bullet(doc, 'Seaborn correlation heatmap')

doc.add_heading('Tab 3: Model Performance', level=2)
add_bullet(doc, 'Comparison table: all 5 models × 5 metrics')
add_bullet(doc, 'Bar chart: AUC-ROC comparison')
add_bullet(doc, 'ROC curves overlay (one per model)')
add_bullet(doc, "Best model's confusion matrix")

doc.add_heading('Tab 4: Predictions & Explainability', level=2)
p = doc.add_paragraph('Interactive inputs:')
p.bold = True
add_bullet(doc, 'st.slider("Credit Limit", 10000, 800000, 150000) — LIMIT_BAL')
add_bullet(doc, 'st.slider("Age", 20, 80, 35) — AGE')
add_bullet(doc, 'st.selectbox("Sex", [1, 2]) — SEX (1=Male, 2=Female)')
add_bullet(doc, 'st.selectbox("Education", [1, 2, 3, 4]) — EDUCATION')
add_bullet(doc, 'st.selectbox("Marriage", [1, 2, 3]) — MARRIAGE')
add_bullet(doc, 'st.slider("Most Recent Payment Status", -2, 8, 0) — PAY_0')
add_bullet(doc, 'Remaining features (BILL_AMT, PAY_AMT) use averaged defaults from training data')
p2 = doc.add_paragraph('Output: ')
p2.add_run('Prediction probability + SHAP waterfall plot for the input').italic = True

doc.add_heading('Tab 5: Ethics & Responsible AI', level=2)
add_bullet(doc, 'Bias analysis results (disparate impact table)')
add_bullet(doc, 'Equalized odds comparison chart')
add_bullet(doc, 'Discussion of societal impact')
add_bullet(doc, 'Mitigation strategies applied/recommended')

# ══════════════════════════════════════════════════════════════
# 9. CV STRATEGY
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. CV Strategy', level=1)
add_code_block(doc, '''from sklearn.model_selection import GridSearchCV, StratifiedKFold

cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)

# Example: Random Forest
param_grid = {
    'n_estimators': [50, 100, 200],
    'max_depth': [3, 5, 8]
}
grid_search = GridSearchCV(
    RandomForestClassifier(random_state=42, class_weight='balanced'),
    param_grid, cv=cv, scoring='roc_auc', n_jobs=-1
)
grid_search.fit(X_train, y_train)''')

# ══════════════════════════════════════════════════════════════
# 10. KAGGLE CLI SETUP
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Appendix: Kaggle CLI Setup', level=1)

doc.add_heading('Step 1: Install Kaggle CLI', level=3)
add_code_block(doc, 'pip install kaggle')

doc.add_heading('Step 2: Set Credentials', level=3)
add_code_block(doc, '''# PowerShell
$env:KAGGLE_USERNAME = "your_kaggle_username"
$env:KAGGLE_KEY = "your_kaggle_api_key"

# Or create %USERPROFILE%\\.kaggle\\kaggle.json:
# {"username": "your_kaggle_username", "key": "your_kaggle_api_key"}''')

doc.add_heading('Step 3: Download Datasets', level=3)
add_code_block(doc, '''# Give Me Some Credit
kaggle competitions download -c GiveMeSomeCredit -p give-me-some-credit/raw/

# Home Credit Default Risk
kaggle competitions download -c home-credit-default-risk -p home-credit-default-risk/raw/

# LendingClub Loans
kaggle datasets download -d ethon0426/lending-club-20072020q1 -p lendingclub-loans/raw/''')

doc.add_heading('Step 4: To obtain your API key', level=3)
add_bullet(doc, 'Go to https://www.kaggle.com/settings')
add_bullet(doc, 'Scroll to "API" section')
add_bullet(doc, 'Click "Create New Token"')
add_bullet(doc, 'Download kaggle.json and place it in %USERPROFILE%\\.kaggle\\')

# ── Save ──
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, 'research_writeup.docx')
doc.save(out_path)
print(f'Saved: {out_path}')

# Also save to OneDrive path
onedrive = r'C:\Users\taash\OneDrive\UW\Class\Winter 2025-2026\MSIS 522 B Advanced Machine Learning\Homework 1 The Complete Data Science Workflow\Option 2 for the credit scoring\research_writeup.docx'
doc.save(onedrive)
print(f'Saved: {onedrive}')
